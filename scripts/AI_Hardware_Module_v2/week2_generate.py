# scripts/week2_generate.py
# 运行说明：
# 1. 确保在主目录 D:\pythonProject\AI_Hardware_Module_v2 下，cd to scripts folder: cd scripts
# 2. 运行: D:\pythonProject\new_venv\Scripts\python.exe week2_generate.py
# 3. 输出: docs/week2_teaching_plan.docx (10页初稿，含文本/表格/截屏/视频链接/代码块)
# 4. 依赖: pip install requests python-docx pillow (Pillow for images, if not installed)
# 5. API: 使用提供的 deepseek_key (hardcoded for testing; replace with file in production)
# 6. 截屏/视频: 自动添加 demos/week2_test_qemu.png 和 YouTube链接 (如果文件缺，用占位符)
# 7. 错误处理: API失败用fallback文本，文件缺用placeholder

import requests
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from PIL import Image  # For image handling if needed
import json  # For quiz JSON

# Step 1: Use provided API key (hardcoded for immediate testing; in production, use file)
api_key = 'sk-ea9c563739364338b71275d7964f37d3'  # Your DeepSeek key from api_keys.txt
print(f"API Key loaded: {api_key[:10]}...")  # Partial print for security

# Optional: Try loading from file if exists (fallback to hardcoded)
config_path = '../config/api_keys.txt'  # Adjusted relative path assuming run from main dir
if os.path.exists(config_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        for line in lines:
            if 'deepseek_key' in line:
                api_key = line.strip().split('=')[1].strip()
                print("Loaded key from file.")
                break

# Step 2: DeepSeek API Prompt for lesson plan text (10 pages content, no dates)
prompt = """
Generate a 10-page lesson plan document for Week 2: Cyber Mechanic - Assembly and Disassembly Practice, based on v2.0 framework. Structure strictly: Cover (title/subtitle/author/version/cover image description); Page 1 (SMART goals + SDT/TAM/CLT path model theory + comparison table with Week 1 connection); Pages 2-5 (timeline table + activity overview + extended details, such as game scoring rules / layered inclusion Lee et al. 5 principles / low-resource optimization); Pages 6-8 (detailed script, including teacher operation list / student interaction examples / full CLIP Demo code snippet / AR QR code generation steps); Pages 9-10 (assessment closed loop, including t-test formula example / SEM β=0.40 simulation / qualitative story gold mine analysis / extended assessment metrics table). Output pure text with clear structure (use [PAGE X] markers for sections), easy to convert to DOCX. Length about 10 pages of content, ensure academic fashion, strong originality (cite real literature like Penchala et al. 2025 ViT heatmap). No dates in content.

Key optimization requirements:
- Connect to Week 1: Start each section with a review of Week 1 circuit exploration (Logisim CPU/memory), such as "Extend from Week 1 farmland brain to disassembly and OS installation", import Week 1 data (interest baseline 3.2 points, logs/week1_results.xlsx).
- Rural localization: Expand the farm machinery disassembly story ("The tractor chip is loose, like Week 1 signal failure"), incorporate CLIP matching ("Chip like farmland screw" matching Week 1 circuit diagram demos/week1_circuit.png).
- Gamification: Detailed scoring rules for the competition (disassembly speed 5 points / CLIP matching 2 points / collaboration 3 points, total 10 points + rewards like "virtual badges").
- Low-resource compatibility: Emphasize QEMU lag fallback (Alpine ISO 200MB + tcg acceleration, or offline video demos/videos/week2_qemu_tutorial.mp4), prioritize mobile AR scanning.
- AI integration: Full CLIP text-image matching code (import transformers/cv2, local path to prevent lag), ViT heatmap generation example (Matplotlib code).
- Inclusion/ethics: Expand Lee et al. 2025 5 principles (cultural relevance / fair access), ethics discussion (Kong & Zhu 2025 CFA pre-post test, data privacy).
- Quantification: Add formula examples, such as t-test (stats.ttest_rel(pre, post) p<0.01), SEM path simulation (autonomy → participation β=0.40).
- Visual/multimedia: Describe attachments (e.g., "Insert demos/week2_collab_heatmap.png, half-page size, caption: Week 2 collaboration heatmap"), video links (YouTube QEMU tutorial).
- Sustainability: Open-source GitHub add Week 2 branch (pull updates), community governance.
"""

# API Call to DeepSeek (detailed error handling)
url = 'https://api.deepseek.com/v1/chat/completions'
headers = {'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}
data = {
    'model': 'deepseek-chat',
    'messages': [{'role': 'user', 'content': prompt}],
    'max_tokens': 6000,  # Increased for detailed 10 pages
    'temperature': 0.7,
    'stream': False
}
try:
    response = requests.post(url, headers=headers, json=data, timeout=60)
    if response.status_code == 200:
        text = response.json()['choices'][0]['message']['content']
        print("API success: Lesson plan text generated.")
    else:
        raise Exception(f"API error {response.status_code}: {response.text}")
except Exception as e:
    print(f"API failed: {e}. Using fallback text.")
    text = """
Fallback Lesson Plan Text (10 pages structure placeholder):
[PAGE 1] SMART Goals: Specific: Master assembly/disassembly. Measurable: Participation +20%. etc.
[PAGE 2] Timeline Table: Time | Activity | Description...
... (expand with details as per prompt)
    """  # Basic fallback

# Step 3: Parse text and build DOCX (detailed structure with tables, code, images)
doc = Document()
doc.add_heading('Week 2 Lesson Plan: Cyber Mechanic - Assembly and Disassembly Practice', 0)
doc.add_page_break()  # Ensure cover on first page

# Cover Page (detailed formatting)
doc.add_heading('Cover', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Title: Week 2: Cyber Mechanic - Assembly and Disassembly Practice').bold = True
p.add_run('\nSubtitle: AI-Driven Low-Resource Rural High School Teaching Framework').bold = True
p.add_run('\nAuthor: Yang Xia')
p.add_run('\nVersion: v2.0')
p.add_run('\nCover Image Description: AR QR code for farm machinery disassembly demo (insert demos/week2_ar_qr.png)')

# Parse sections by [PAGE X] markers (detailed splitting)
pages = {}
if '[PAGE 1]' in text:
    pages = {i: text.split(f'[PAGE {i}]')[1].split(f'[PAGE {i+1}]')[0] if f'[PAGE {i+1}]' in text else text.split(f'[PAGE {i}]')[1] for i in range(1, 11)}
else:
    # Fallback split by length
    sections = text.split('\n\n')
    for i in range(1, 11):
        pages[i] = '\n'.join(sections[(i-1)*2:i*2]) if len(sections) >= i*2 else 'Placeholder content for Page ' + str(i)

# Page 1: Goals, Theory, Week 1 Connection (detailed table)
doc.add_page_break()
doc.add_heading('Page 1: SMART Goals + Theory + Week 1 Connection', level=1)
doc.add_paragraph(pages[1])  # Add text

# Detailed Week 1 Connection Table
table = doc.add_table(rows=5, cols=4)  # 5 rows for detail
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Aspect'
hdr_cells[1].text = 'Week 1 (Circuit Exploration)'
hdr_cells[2].text = 'Week 2 (Cyber Mechanic)'
hdr_cells[3].text = 'Connection'
# Row 1
row1 = table.rows[1].cells
row1[0].text = 'Core Tool'
row1[1].text = 'Logisim simulation'
row1[2].text = 'QEMU virtual disassembly'
row1[3].text = 'Extend circuit to OS installation'
# Row 2
row2 = table.rows[2].cells
row2[0].text = 'Knowledge Point'
row2[1].text = 'CPU/Memory concepts'
row2[2].text = 'Chip installation/OS preload'
row2[3].text = 'From signal to software run'
# Row 3
row3 = table.rows[3].cells
row3[0].text = 'Quantitative Goal'
row3[1].text = 'Interest +15% (baseline 3.2)'
row3[2].text = 'Participation +20% (ViT ≥85%)'
row3[3].text = 'Cumulative literacy +25%'
# Row 4
row4 = table.rows[4].cells
row4[0].text = 'Story Element'
row4[1].text = 'Farmland signal transmission'
row4[2].text = 'Farm machinery chip disassembly'
row4[3].text = 'Extend farmland brain to tractor repair'

# Pages 2-5: Timeline, Activities, Details (detailed tables)
for i in range(2, 6):
    doc.add_page_break()
    doc.add_heading(f'Page {i}: Timeline & Activities', level=1)
    doc.add_paragraph(pages[i])
    # Detailed Timeline Table (5 columns, 5 rows example)
    timeline_table = doc.add_table(rows=5, cols=5)
    timeline_table.style = 'Table Grid'
    hdr = timeline_table.rows[0].cells
    hdr[0].text = 'Time Segment'
    hdr[1].text = 'Activity Link'
    hdr[2].text = 'Detailed Description'
    hdr[3].text = 'Tools/Resources'
    hdr[4].text = 'Teacher Role/Student Output'
    # Fill example rows (expand with text)
    row1 = timeline_table.rows[1].cells
    row1[0].text = '0-10min'
    row1[1].text = 'Introduction & Activation'
    row1[2].text = 'Review Week 1 story, introduce farm machinery disassembly'
    row1[3].text = 'DeepSeek story script, MagicSchool quiz'
    row1[4].text = 'Guide review; Student brain storm screenshots'
    # Add 3 more rows similarly (parse from pages[i] or extend)
    for j in range(2, 5):
        row = timeline_table.add_row().cells
        row[0].text = f'{10*(j-1)}-{10*j}min' if j < 4 else '40-45min'
        row[1].text = ['Practice Exploration', 'Sharing & Assessment', 'Reflection & Extension'][j-2]
        row[2].text = ['Virtual disassembly with QEMU', 'Group demo and quiz', 'Ethics and global trends'][j-2]
        row[3].text = ['QEMU, CLIP code', 'Grok quiz, Matplotlib', 'Character.AI, Notion'][j-2]
        row[4].text = ['Supervise; Student screenshots', 'Facilitate; Student scores', 'Summarize; Student logs'][j-2]

# Pages 6-8: Script, CLIP Code, AR Steps (detailed code block)
for i in range(6, 9):
    doc.add_page_break()
    doc.add_heading(f'Page {i}: Detailed Script & Demos', level=1)
    doc.add_paragraph(pages[i])
    # Full CLIP Code Block (detailed)
    doc.add_paragraph('Full CLIP Text-Image Matching Code:')
    code_para = doc.add_paragraph()
    code_run = code_para.add_run()
    code_run.text = 'from transformers import CLIPProcessor, CLIPModel\n'
    code_run.text += 'import cv2\n'
    code_run.text += 'model = CLIPModel.from_pretrained("./clip-vit-base-patch32")  # Local path to prevent lag\n'
    code_run.text += 'processor = CLIPProcessor.from_pretrained("./clip-vit-base-patch32")\n'
    code_run.text += 'image = cv2.imread("../demos/week1_circuit.png")  # Connect to Week 1 circuit\n'
    code_run.text += 'inputs = processor(text=["Chip like farmland screw"], images=image, return_tensors="pt")\n'
    code_run.text += 'outputs = model(**inputs)\n'
    code_run.text += 'print("Matching score:", outputs.logits_per_image)  # >0.8 success'
    code_run.font.name = 'Courier New'  # Monospace for code
    code_run.font.size = Inches(0.1)  # Small font
    # AR QR Code Generation Steps (list)
    doc.add_paragraph('AR QR Code Generation Steps:')
    doc.add_paragraph('1. Open Tinkercad.com, create farm machinery model.\n2. Export > AR > Download QR PNG to demos/week2_ar_qr.png.\n3. Mobile scan for disassembly demo.')

# Pages 9-10: Assessment (detailed formulas/table)
for i in range(9, 11):
    doc.add_page_break()
    doc.add_heading(f'Page {i}: Assessment Closed Loop', level=1)
    doc.add_paragraph(pages[i])
    # t-test Formula Example
    doc.add_paragraph('t-test Formula Example:')
    code_para = doc.add_paragraph()
    code_run = code_para.add_run()
    code_run.text = 'from scipy.stats import ttest_rel\npre = [3.2, 3.1, 3.3]  # Week 1 baseline data\npost = [3.7, 3.6, 3.8]  # Week 2 post-test\nt_stat, p_val = ttest_rel(pre, post)\nprint(f"t={t_stat}, p={p_val}")  # Expected p<0.01'
    code_run.font.name = 'Courier New'
    # Extended Assessment Metrics Table
    metrics_table = doc.add_table(rows=4, cols=3)
    metrics_table.style = 'Table Grid'
    hdr = metrics_table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Target Value'
    hdr[2].text = 'Measurement Tool'
    # Row 1
    row1 = metrics_table.rows[1].cells
    row1[0].text = 'Interest Improvement'
    row1[1].text = '+15% (3.7 score)'
    row1[2].text = 'Likert pre/post Excel'
    # Row 2
    row2 = metrics_table.rows[2].cells
    row2[0].text = 'Participation Rate'
    row2[1].text = '≥85%'
    row2[2].text = 'ViT heatmap (97.58% accuracy)'
    # Row 3
    row3 = metrics_table.rows[3].cells
    row3[0].text = 'Quiz Accuracy'
    row3[1].text = '85%'
    row3[2].text = 'JSON quiz scores'

# Step 4: Add Screenshots and Video Links (detailed image handling)
doc.add_page_break()
doc.add_heading('Attachments: Screenshots and Video Links', level=1)
# Add QEMU Screenshot (check if exists)
qemu_img_path = '../demos/week2_test_qemu.png'
if os.path.exists(qemu_img_path):
    try:
        p = doc.add_paragraph('QEMU Virtual Machine Screenshot (Week 2 Demo):')
        run = p.add_run()
        run.add_picture(qemu_img_path, width=Inches(5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("Added QEMU screenshot to DOCX.")
    except Exception as e:
        doc.add_paragraph(f'Placeholder: Insert {qemu_img_path} (QEMU startup interface, half-page size, caption: Week 2 virtual disassembly demo)')
        print(f"Image add error: {e}, used placeholder.")
else:
    doc.add_paragraph('Placeholder: Add demos/week2_test_qemu.png after running stage 1.')

# Add Video Link (hyperlink)
doc.add_paragraph('YouTube QEMU Tutorial Video Link:')
video_para = doc.add_paragraph()
hyperlink = video_para.add_hyperlink('QEMU Ubuntu Installation Tutorial', 'https://www.youtube.com/watch?v=Oz-e4atTFkQ')
hyperlink.font.color.rgb = RGBColor(0, 0, 255)  # Blue link
video_para.add_paragraph('Description: 5-min offline fallback video for low-resource demo (demos/videos/week2_qemu_tutorial.mp4).')

# Step 5: Generate Quiz JSON (bonus, detailed)
prompt_quiz = "Generate Week 2 quiz: 3 questions on OS disassembly + CLIP matching, JSON format, adaptive difficulty, connect to Week 1 CPU knowledge."
try:
    response_quiz = requests.post(url, headers=headers, json={'model': 'deepseek-chat', 'messages': [{'role': 'user', 'content': prompt_quiz}], 'max_tokens': 500}, timeout=30)
    if response_quiz.status_code == 200:
        quiz_text = response_quiz.json()['choices'][0]['message']['content']
        quiz_path = '../quiz_library/week2_quiz.json'
        os.makedirs(os.path.dirname(quiz_path), exist_ok=True)
        with open(quiz_path, 'w', encoding='utf-8') as f:
            f.write(quiz_text)
        print(f"Generated {quiz_path}")
    else:
        print("Quiz API error, skipped.")
except Exception as e:
    print(f"Quiz generation error: {e}")

# Step 6: Save DOCX (detailed output)
output_path = '../docs/week2_teaching_plan.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Generated {output_path} (10-page draft with tables, code, screenshots, and video links). Open in Word to review/edit. Total pages: {len([p for p in doc.paragraphs if p.text.strip()])} sections.")

# Optional: Add ViT Heatmap Code to Appendix
doc.add_page_break()
doc.add_heading('Appendix: ViT Heatmap Generation Example', level=1)
code_para = doc.add_paragraph()
code_run = code_para.add_run()
code_run.text = 'import matplotlib.pyplot as plt\nimport numpy as np\nheatmap_data = np.random.rand(25, 25)  # 25 students collaboration\nplt.figure(figsize=(6, 6))\nplt.imshow(heatmap_data, cmap="hot")\nplt.title("Week 2 Collaboration Heatmap")\nplt.savefig("../demos/week2_collab_heatmap.png")\nplt.show()'
code_run.font.name = 'Courier New'